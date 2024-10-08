from docx import Document
import pandas as pd
from docx.shared import Pt
from docx.oxml import OxmlElement, ns
from docx2pdf import convert
import pythoncom
from concurrent.futures import ThreadPoolExecutor, as_completed

# Função para gerar documentos com paralelismo na conversão para PDF
def generate_documents(df, template_paths):
    documentos_gerados = []
    
    # ThreadPoolExecutor para gerar documentos DOCX
    with ThreadPoolExecutor() as doc_executor:
        futures = []
        
        # Para cada modelo selecionado, adicionar a função de geração ao executor
        for template_path in template_paths:
            if template_path == "modelo_r1.docx":
                futures.append(doc_executor.submit(generate_documents_r1, df, documentos_gerados))
            elif template_path == "modelo_r2.docx":
                futures.append(doc_executor.submit(generate_documents_r2, df, documentos_gerados))
            elif template_path == "modelo_da.docx":
                futures.append(doc_executor.submit(generate_documents_da, df, documentos_gerados))
        
        # Aguarda todas as threads terminarem a geração de documentos DOCX
        for future in as_completed(futures):
            future.result()  # Assegura que exceções são capturadas

    # Depois de gerar os documentos DOCX, cria um novo executor para converter PDFs
    pdf_conversion_tasks = []
    with ThreadPoolExecutor() as pdf_executor:
        # Executa a conversão para PDF de forma sequencial ou em paralelo
        for docx_file in documentos_gerados:
            if docx_file.endswith(".docx"):
                pdf_conversion_tasks.append(pdf_executor.submit(convert_pdf_sequential, docx_file))

        # Aguarda todas as conversões para PDF terminarem
        for future in as_completed(pdf_conversion_tasks):
            future.result()

    return documentos_gerados

def convert_pdf_sequential(output_file_docx):
    pythoncom.CoInitialize()
    try:
        output_file_pdf = output_file_docx.replace('.docx', '.pdf')
        convert(output_file_docx, output_file_pdf)
    finally:
        pythoncom.CoUninitialize()

def generate_documents_r1(df, documentos_gerados):
    # Agrupar membros do colegiado e votos por Número do Auto
    df_grouped = df.groupby('NumeroAuto').agg({
        'NumeroProcesso': 'first',
        'ResultadoJulgamento': 'first',
        'DataJulgamento': 'first',
        'Recorrente': 'first',
        'CPF_CNPJ': 'first',
        'AlegacaoRecorrente': 'first',
        'Fundamentacao': 'first',
        'Colegiado': list,  # Usamos list para depois formatar corretamente
        'Voto': list        # Usamos list para depois formatar corretamente
    }).reset_index()

    # Itera sobre os autos agrupados e preenche o documento
    for index, row in df_grouped.iterrows():
        # Carrega o documento do template
        doc = Document("templates/modelo_r1.docx")

        # Remover qualquer tabela existente no template (se houver)
        for table in doc.tables:
            table._element.getparent().remove(table._element)

        # Preenche os outros campos do documento manualmente
        for paragraph in doc.paragraphs:
            if '{{NumeroAuto}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{NumeroAuto}}', str(row['NumeroAuto']))
            if '{{NumeroProcesso}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{NumeroProcesso}}', str(row['NumeroProcesso']))
            if '{{ResultadoJulgamento}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{ResultadoJulgamento}}', str(row['ResultadoJulgamento']))
            if '{{DataJulgamento}}' in paragraph.text:
                 # Formata a data no formato brasileiro
                formatted_date = formatar_data_br(row['DataJulgamento'])
                paragraph.text = paragraph.text.replace('{{DataJulgamento}}', formatted_date)
            if '{{Recorrente}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{Recorrente}}', str(row['Recorrente']))
            if '{{CPF_CNPJ}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{CPF_CNPJ}}', str(row['CPF_CNPJ']))
            if '{{AlegacaoRecorrente}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{AlegacaoRecorrente}}', str(row['AlegacaoRecorrente']))
            if '{{Fundamentacao}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{Fundamentacao}}', str(row['Fundamentacao']))

        # Procurar o parágrafo que contém "Nº DA SESSÃO DE JULGAMENTO"
        for i, paragraph in enumerate(doc.paragraphs):
            if "Nº DA SESSÃO DE JULGAMENTO" in paragraph.text:
                # Adiciona a tabela logo após esse parágrafo
                insert_position = i + 1
                break

        # Adiciona uma nova tabela com a quantidade de membros do colegiado e votos
        table = doc.add_table(rows=1, cols=2)  # Cabeçalhos de duas colunas

        # Definir os títulos das colunas
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'MEMBRO DO COLEGIADO'
        hdr_cells[1].text = 'VOTO'

        # Adicionar os dados do colegiado e voto na tabela
        for colegiado, voto in zip(row['Colegiado'], row['Voto']):
            row_cells = table.add_row().cells
            row_cells[0].text = colegiado  # Coluna de Colegiado
            row_cells[1].text = voto  # Coluna de Voto

        # Adiciona bordas à tabela
        set_table_borders(table)
        doc.paragraphs[insert_position]._element.addnext(table._element)

        # Salva o documento preenchido
        output_file_docx = f"data/INTEIRO TEOR RECURSO 1ª INSTÂNCIA {row['NumeroAuto']}.docx"
        doc.save(output_file_docx)
        documentos_gerados.append(output_file_docx)

def generate_documents_r2(df, documentos_gerados):
    # Agrupar membros do colegiado e votos por Número do Auto
    df_grouped = df.groupby('NumeroAuto').agg({
        'NumeroProcesso': 'first',
        'ResultadoJulgamento': 'first',
        'DataJulgamento': 'first',
        'Recorrente': 'first',
        'CPF_CNPJ': 'first',
        'AlegacaoRecorrente': 'first',
        'Fundamentacao': 'first',
        'Colegiado': list,  # Usamos list para depois formatar corretamente
        'Voto': list        # Usamos list para depois formatar corretamente
    }).reset_index()

    # Itera sobre os autos agrupados e preenche o documento
    for index, row in df_grouped.iterrows():
        # Carrega o documento do template
        doc = Document("templates/modelo_r2.docx")

        # Remover qualquer tabela existente no template (se houver)
        for table in doc.tables:
            table._element.getparent().remove(table._element)

        # Preenche os outros campos do documento manualmente
        for paragraph in doc.paragraphs:
            if '{{NumeroAuto}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{NumeroAuto}}', str(row['NumeroAuto']))
            if '{{NumeroProcesso}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{NumeroProcesso}}', str(row['NumeroProcesso']))
            if '{{ResultadoJulgamento}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{ResultadoJulgamento}}', str(row['ResultadoJulgamento']))
            if '{{DataJulgamento}}' in paragraph.text:
                 # Formata a data no formato brasileiro
                formatted_date = formatar_data_br(row['DataJulgamento'])
                paragraph.text = paragraph.text.replace('{{DataJulgamento}}', formatted_date)
            if '{{Recorrente}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{Recorrente}}', str(row['Recorrente']))
            if '{{CPF_CNPJ}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{CPF_CNPJ}}', str(row['CPF_CNPJ']))
            if '{{AlegacaoRecorrente}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{AlegacaoRecorrente}}', str(row['AlegacaoRecorrente']))
            if '{{Fundamentacao}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{Fundamentacao}}', str(row['Fundamentacao']))

        # Procurar o parágrafo que contém "Nº DA SESSÃO DE JULGAMENTO"
        for i, paragraph in enumerate(doc.paragraphs):
            if "Nº DA SESSÃO DE JULGAMENTO" in paragraph.text:
                # Adiciona a tabela logo após esse parágrafo
                insert_position = i + 1
                break

        # Adiciona uma nova tabela com a quantidade de membros do colegiado e votos
        table = doc.add_table(rows=1, cols=2)  # Cabeçalhos de duas colunas

        # Definir os títulos das colunas
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'MEMBRO DO COLEGIADO'
        hdr_cells[1].text = 'VOTO'

        # Adicionar os dados do colegiado e voto na tabela
        for colegiado, voto in zip(row['Colegiado'], row['Voto']):
            row_cells = table.add_row().cells
            row_cells[0].text = colegiado  # Coluna de Colegiado
            row_cells[1].text = voto  # Coluna de Voto

        # Adiciona bordas à tabela
        set_table_borders(table)

        # Insere a tabela no local correto (após "Nº DA SESSÃO DE JULGAMENTO")
        doc.paragraphs[insert_position]._element.addnext(table._element)

        # Salva o documento preenchido
        output_file_docx = f"data/INTEIRO TEOR RECURSO 2ª INSTÂNCIA {row['NumeroAuto']}.docx"
        doc.save(output_file_docx)
        documentos_gerados.append(output_file_docx)

def generate_documents_da(df, documentos_gerados):
    # Agrupar membros do colegiado e votos por Número do Auto
    df_grouped = df.groupby('NumeroAuto').agg({
        'NumeroProcesso': 'first',
        'ResultadoJulgamento': 'first',
        'DataJulgamento': 'first',
        'Recorrente': 'first',
        'CPF_CNPJ': 'first',
        'AlegacaoRecorrente': 'first',
        'Fundamentacao': 'first',
    }).reset_index()

    # Itera sobre os autos agrupados e preenche o documento
    for index, row in df_grouped.iterrows():
        # Carrega o documento do template
        doc = Document("templates/modelo_da.docx")

        # Preenche os outros campos do documento manualmente
        for paragraph in doc.paragraphs:
            if '{{NumeroAuto}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{NumeroAuto}}', str(row['NumeroAuto']))
            if '{{NumeroProcesso}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{NumeroProcesso}}', str(row['NumeroProcesso']))
            if '{{ResultadoJulgamento}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{ResultadoJulgamento}}', str(row['ResultadoJulgamento']))
            if '{{DataJulgamento}}' in paragraph.text:
                # Formata a data no formato brasileiro
                formatted_date = formatar_data_br(row['DataJulgamento'])
                paragraph.text = paragraph.text.replace('{{DataJulgamento}}', formatted_date)
            if '{{Recorrente}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{Recorrente}}', str(row['Recorrente']))
            if '{{CPF_CNPJ}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{CPF_CNPJ}}', str(row['CPF_CNPJ']))
            if '{{AlegacaoRecorrente}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{AlegacaoRecorrente}}', str(row['AlegacaoRecorrente']))
            if '{{Fundamentacao}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{Fundamentacao}}', str(row['Fundamentacao']))

        # Salva o documento preenchido
        output_file_docx = f"data/INTEIRO TEOR DEFESA DA AUTUAÇÃO {row['NumeroAuto']}.docx"
        doc.save(output_file_docx)
        documentos_gerados.append(output_file_docx)

def set_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')

            # Define bordas para todas as direções
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(ns.qn('w:val'), 'single')  # Atributo correto para o namespace
                border.set(ns.qn('w:sz'), '4')  # Tamanho da borda
                border.set(ns.qn('w:space'), '0')
                border.set(ns.qn('w:color'), '000000')  # Cor preta
                tcBorders.append(border)
            
            tcPr.append(tcBorders)

# Função para formatar a data no formato brasileiro
def formatar_data_br(data):
    return pd.to_datetime(data).strftime('%d/%m/%Y')

def convert_pdf(output_file_docx, pdfs_gerados):
    pythoncom.CoInitialize()

    try:
        output_file_pdf = output_file_docx.replace('.docx', '.pdf')
        
        # Converte o documento DOCX para PDF usando docx2pdf
        convert(output_file_docx, output_file_pdf)
        
        pdfs_gerados.append(output_file_pdf)

    finally:
        pythoncom.CoUninitialize()
