import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement, ns  # Importa também o namespace

# Função para adicionar bordas às células da tabela
def set_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')

            # Define bordas para todas as direções
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(ns.qn('w:val'), 'single')  # Atributo correto para o namespace
                border.set(ns.qn('w:sz'), '4')  # Tamanho da borda
                border.set(ns.qn('w:space'), '0')
                border.set(ns.qn('w:color'), '000000')  # Cor preta
                tcBorders.append(border)
            
            tcPr.append(tcBorders)

# Carrega os dados da planilha
df = pd.read_excel('dados.xlsx')

# Agrupar membros do colegiado e votos por Número do Auto
df_grouped = df.groupby('NumeroAuto').agg({
    'NumeroProcesso': 'first',
    'ResultadoJulgamento': 'first',
    'DataJulgamento': 'first',
    'Recorrente': 'first',
    'CPF_CNPJ': 'first',
    'AlegacaoRecorrente': 'first',
    'Fundamentacao': 'first',
    'Colegiado': list,  # Usamos list para depois formatar corretamente
    'Voto': list        # Usamos list para depois formatar corretamente
}).reset_index()

# Itera sobre os autos agrupados e preenche o documento
for index, row in df_grouped.iterrows():
    # Carrega o documento do template
    doc = Document("template.docx")

    # Remover qualquer tabela existente no template (se houver)
    for table in doc.tables:
        table._element.getparent().remove(table._element)

    # Preenche os outros campos do documento manualmente
    for paragraph in doc.paragraphs:
        if '{{NumeroAuto}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{NumeroAuto}}', str(row['NumeroAuto']))
        if '{{NumeroProcesso}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{NumeroProcesso}}', str(row['NumeroProcesso']))
        if '{{ResultadoJulgamento}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ResultadoJulgamento}}', str(row['ResultadoJulgamento']))
        if '{{DataJulgamento}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{DataJulgamento}}', str(row['DataJulgamento']))
        if '{{Recorrente}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{Recorrente}}', str(row['Recorrente']))
        if '{{CPF_CNPJ}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{CPF_CNPJ}}', str(row['CPF_CNPJ']))
        if '{{AlegacaoRecorrente}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{AlegacaoRecorrente}}', str(row['AlegacaoRecorrente']))
        if '{{Fundamentacao}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{Fundamentacao}}', str(row['Fundamentacao']))

    # Procurar o parágrafo que contém "Nº DA SESSÃO DE JULGAMENTO"
    for i, paragraph in enumerate(doc.paragraphs):
        if "Nº DA SESSÃO DE JULGAMENTO" in paragraph.text:
            # Adiciona a tabela logo após esse parágrafo
            insert_position = i + 1
            break

    # Adiciona uma nova tabela com a quantidade de membros do colegiado e votos
    table = doc.add_table(rows=1, cols=2)  # Cabeçalhos de duas colunas

    # Definir os títulos das colunas
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'MEMBRO DO COLEGIADO'
    hdr_cells[1].text = 'VOTO'

    # Adicionar os dados do colegiado e voto na tabela
    for colegiado, voto in zip(row['Colegiado'], row['Voto']):
        row_cells = table.add_row().cells
        row_cells[0].text = colegiado  # Coluna de Colegiado
        row_cells[1].text = voto  # Coluna de Voto

    # Adiciona bordas à tabela
    set_table_borders(table)

    # Insere a tabela no local correto (após "Nº DA SESSÃO DE JULGAMENTO")
    doc.paragraphs[insert_position]._element.addnext(table._element)

    # Salva o documento preenchido
    output_file = f"output_{row['NumeroAuto']}.docx"
    doc.save(output_file)

    print(f"Documento gerado: {output_file}")
